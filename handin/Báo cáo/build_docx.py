"""
Chuyển báo cáo NoCNet-v2 (Markdown) -> Word (.docx) theo mẫu báo cáo.
Tự nhúng sơ đồ kiến trúc (thay khối mermaid) và các hình kết quả vào đúng mục.

Chạy:  python build_docx.py
Đầu ra:  BaoCao_NoCNetV2_VI.docx
"""
import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(HERE, "BaoCao_NoCNetV2_VI.md")
ARCH = os.path.join(HERE, "nocnet_v2_arch.png")
RES = os.path.normpath(os.path.join(HERE, "..", "Kết quả"))

NAVY = RGBColor(0x1F, 0x4E, 0x79)
FONT = "Times New Roman"
fig_counter = [0]


def set_default_font(doc, name=FONT, size=13):
    st = doc.styles["Normal"]
    st.font.name = name
    st.font.size = Pt(size)
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), name)


def clean_math(s):
    s = s.replace("$$", "").replace("$", "")
    s = re.sub(r"\\mathcal\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\text\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\mathrm\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\frac\{([^}]*)\}\{([^}]*)\}", r"(\1)/(\2)", s)
    s = re.sub(r"\\sqrt\{([^}]*)\}", r"√(\1)", s)
    s = s.replace("\\cdot", "·").replace("\\times", "×").replace("\\top", "ᵀ")
    s = s.replace("\\Sigma", "Σ").replace("\\sum", "Σ").replace("\\Delta", "Δ")
    s = s.replace("\\le", "≤").replace("\\ge", "≥").replace("\\approx", "≈")
    s = s.replace("\\gamma", "γ").replace("\\alpha", "α").replace("\\beta", "β").replace("\\rho", "ρ").replace("\\phi", "φ")
    s = s.replace("\\,", " ").replace("\\;", " ").replace("\\quad", "  ")
    s = re.sub(r"_\{([^}]*)\}", r"(\1)", s)
    s = re.sub(r"\^\\top", "ᵀ", s)
    s = re.sub(r"\^\{([^}]*)\}", r"^\1", s)
    s = s.replace("\\", "")
    return s


INLINE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*|`[^`]+?`)")


def add_runs(p, text):
    text = re.sub(r"\[\[([^\]]+)\]\]", r"\1", text)
    text = re.sub(r"\$([^$]+)\$", lambda m: clean_math(m.group(1)), text)
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(11)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = p.add_run(part[1:-1]); r.italic = True
        else:
            p.add_run(part)


def add_image(doc, path, caption, width=6.0):
    if not os.path.exists(path):
        print("  ! thiếu hình:", path); return
    fig_counter[0] += 1
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(f"Hình {fig_counter[0]}. {caption}"); r.italic = True; r.font.size = Pt(11)


def style_table(table):
    table.style = "Light Grid Accent 1"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)


def build():
    lines = open(MD, encoding="utf-8").read().split("\n")
    doc = Document()
    set_default_font(doc)
    for sec in doc.sections:
        sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]; stripped = line.strip()

        if stripped.startswith("```"):
            lang = stripped[3:].strip().lower()
            j = i + 1; buf = []
            while j < n and not lines[j].strip().startswith("```"):
                buf.append(lines[j]); j += 1
            if lang == "mermaid":
                add_image(doc, ARCH, "Sơ đồ kiến trúc tổng thể (top-down) của mô hình "
                          "đề xuất NoCNet-v2 (Deep Sets).")
            else:
                p = doc.add_paragraph(); r = p.add_run("\n".join(buf))
                r.font.name = "Consolas"; r.font.size = Pt(10)
            i = j + 1; continue

        if stripped in ("---", "***", "___"):
            i += 1; continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1)); txt = m.group(2).strip()
            if level == 1:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(clean_math(re.sub(r"\*\*", "", txt)))
                r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
            else:
                h = doc.add_heading(level=min(level - 1, 4)); add_runs(h, txt)
                for r in h.runs:
                    r.font.color.rgb = NAVY
            i += 1; continue

        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|[\s:|-]+\|?\s*$", lines[i + 1].strip()):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            rows = []; j = i + 2
            while j < n and lines[j].strip().startswith("|"):
                rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")]); j += 1
            tbl = doc.add_table(rows=1, cols=len(header))
            for k, htext in enumerate(header):
                cell = tbl.rows[0].cells[k]; cell.paragraphs[0].text = ""
                add_runs(cell.paragraphs[0], htext)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            for row in rows:
                cells = tbl.add_row().cells
                for k in range(len(header)):
                    val = row[k] if k < len(row) else ""
                    cells[k].paragraphs[0].text = ""; add_runs(cells[k].paragraphs[0], val)
            style_table(tbl)
            h0 = header[0].lower()
            if h0.startswith("độ đo"):
                add_image(doc, os.path.join(RES, "confusion_matrix.png"),
                          "Ma trận nhầm lẫn trên tập kiểm thử (923 profile, grouped seed 42).", 5.4)
                add_image(doc, os.path.join(RES, "training_history.png"),
                          "Lịch sử huấn luyện (tinh chỉnh trên PROVEDIt thật).", 6.0)
            elif h0 == "noc" and any("accuracy" in c.lower() for c in header):
                add_image(doc, os.path.join(RES, "per_class.png"),
                          "Accuracy và F1 theo từng lớp NoC.", 5.6)
            elif h0.startswith("hệ thống"):
                add_image(doc, os.path.join(RES, "system_comparison.png"),
                          "So sánh độ chính xác đếm NoC giữa các hệ trên PROVEDIt.", 6.0)
            i = j; continue

        if stripped.startswith(">"):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, stripped.lstrip(">").strip())
            for r in p.runs:
                r.italic = True
            i += 1; continue

        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet"); add_runs(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1; continue

        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number"); add_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1; continue

        if stripped.startswith("$$"):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(clean_math(stripped)); r.italic = True
            i += 1; continue

        if not stripped:
            i += 1; continue

        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
        add_runs(p, stripped)
        i += 1

    out = os.path.join(HERE, "BaoCao_NoCNetV2_VI.docx")
    doc.save(out)
    print("Đã tạo:", out, "| số hình nhúng:", fig_counter[0])


if __name__ == "__main__":
    build()
