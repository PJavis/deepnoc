"""
Sinh báo cáo NoCNet-v2 bằng tiếng Việt (.docx).

Cấu trúc:
  1. Tóm tắt
  2. Bối cảnh: bài toán NoC + paper deepNoC + repo hiện tại
  3. Chẩn đoán: vì sao baseline (NoCFormer, deepNoC) chưa đủ
  4. Kiến trúc NoCNet-v2 (Deep Sets per-locus, stutter-bias attention,
     cross-locus Transformer, count-aware multi-head)
  5. Dữ liệu tổng hợp: superposition + physics (stutter regen, dropout, noise)
  6. Loss đa nhiệm
  7. Pipeline huấn luyện: hybrid loader, SWA, fine-tune
  8. Đánh giá: 5-fold grouped CV
  9. Kết luận + roadmap

Chạy:
    python report/build_nocnet_v2_report.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# -------------------- Style helpers --------------------------------------

def set_default_font(doc: Document, name: str = "Calibri", size_pt: int = 11):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)
    rpr = style.element.rPr
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x2D, 0x5C)
    return h


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)


def add_code(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    rpr = r._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")
    rpr.append(rfonts)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)


def add_table(doc: Document, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val
    return table


def add_image(doc: Document, path: str, width_in: float = 6.5, caption: str = ""):
    if not os.path.exists(path):
        add_para(doc, f"[Không tìm thấy ảnh: {path}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)


# -------------------- Build report ---------------------------------------

def build():
    here = os.path.dirname(__file__)
    doc = Document()
    set_default_font(doc)

    # ==================== TRANG TIÊU ĐỀ ====================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("NoCNet-v2")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1F, 0x2D, 0x5C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Deep Sets + stutter-aware attention + count-aware head "
                    "cho bài toán Number of Contributors")
    r.italic = True
    r.font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Báo cáo cải tiến tiếp theo cho deepNoC — "
                 "Taylor & Humphries (2024), arXiv:2412.09803\n").italic = True
    meta.add_run("Repo: deepNoC (kiến trúc nocnet_v2 mới)\n")
    meta.add_run("Trạng thái: code + smoke test xong; "
                 "đợi train full trên GPU 16 GB\n")

    doc.add_paragraph()

    # ==================== 1. TÓM TẮT ====================
    add_heading(doc, "1. Tóm tắt", level=1)
    add_para(doc,
        "Báo cáo này trình bày NoCNet-v2 — kiến trúc mới được thiết kế riêng cho "
        "bài toán dự đoán số người đóng góp DNA (NoC) trên bộ PROVEDIt. "
        "Mục tiêu là vượt cả paper gốc deepNoC (Taylor & Humphries, 2024) lẫn "
        "phiên bản NoCFormer trước đó trong repo. So với NoCFormer, NoCNet-v2 "
        "thay đổi cả ba thành phần lớn của pipeline: kiến trúc model, cách sinh "
        "dữ liệu, và cơ chế đánh giá.")

    add_para(doc, "Bốn cải tiến chính:", bold=True)
    add_bullet(doc, "(1) Pooling kiểu Deep Sets [sum, max, log1p(count)] thay "
                    "cho CLS-token attention — giữ trực tiếp tín hiệu \"đếm\" "
                    "vốn là cốt lõi của bài toán NoC.")
    add_bullet(doc, "(2) Self-attention nội locus có additive bias theo "
                    "allele-distance Δ(a_i, a_j), thay vì stutter feature "
                    "thủ công — mô hình tự học cấu trúc stutter (±1, ±2, "
                    "±0.2 repeat).")
    add_bullet(doc, "(3) Synthetic mixture pool có physics đầy đủ: "
                    "superposition + stutter regen + allelic dropout + "
                    "drop-in noise + LOD filter, lấy từ pool ~2700 profile "
                    "NoC=1 của PROVEDIt.")
    add_bullet(doc, "(4) Đầu ra 3-view (softmax + scalar regression + CORN "
                    "ordinal) được ensemble ở inference; SWA tự kích hoạt "
                    "trong 25% epoch cuối; fine-tune real-data có "
                    "baseline-guard chống regression.")

    add_para(doc,
        "Smoke test xong toàn bộ pipeline trên GPU: model 574k params, "
        "peak VRAM ~1.3 GiB ở batch 32 — vừa GPU 6 GB cá nhân. "
        "Kỳ vọng hợp lý sau khi train full trên GPU 16 GB: 80–85% accuracy "
        "honest (grouped split, không leak replicate). Mục tiêu 90% còn cần "
        "thêm ensemble nhiều seed + tune threshold per-class.")

    # ==================== 2. BỐI CẢNH ====================
    add_heading(doc, "2. Bối cảnh", level=1)
    add_heading(doc, "2.1 Bài toán NoC", level=2)
    add_para(doc,
        "Number of Contributors (NoC) là số người có DNA đóng góp vào một "
        "mẫu STR thu được. Mỗi mẫu được phân tích trên kit STR (ở đây là "
        "GlobalFiler — 24 locus), kết quả là electropherogram thể hiện các "
        "peak ở từng locus. Việc xác định đúng NoC quyết định mọi phân tích "
        "kế tiếp (deconvolution, likelihood ratio…). Phương pháp thủ công "
        "phổ biến nhất là MAC (Maximum Allele Count), nhưng dưới-đếm khi có "
        "allele sharing giữa các contributor.")

    add_heading(doc, "2.2 Dữ liệu PROVEDIt và bộ filter", level=2)
    add_para(doc,
        "Repo dùng phần PROVEDIt 1-5 person CSVs đã filter "
        "(`GlobalFiler + ABI 3500 + 25 sec`). Sau bước `prepare`:")
    add_bullet(doc, "Tổng số profile: 3378")
    add_bullet(doc, "NoC=1: 2712 (rất nhiều) — nguồn cho synthetic pool")
    add_bullet(doc, "NoC=2..5: 174 / 160 / 176 / 156 (rất ít)")
    add_para(doc,
        "Mỗi profile được chuyển thành tensor `[24 × 50 × 89]`: 24 locus, "
        "tối đa 50 peak mỗi locus, 89 đặc trưng cho mỗi peak (one-hot locus, "
        "allele, size, height, allele freq, peak label probability, các "
        "feature stutter, đếm peak ở mức locus/profile, mixture proportion).")
    layout_png = os.path.join(here, "dataset_layout.png")
    add_image(doc, layout_png, width_in=6.0,
              caption="Hình 1. Layout dữ liệu PROVEDIt: profile → 24 locus × "
                      "50 peak × 89 feature.")

    add_heading(doc, "2.3 Paper deepNoC và NoCFormer", level=2)
    add_para(doc,
        "Paper gốc của Taylor & Humphries (2024) đề xuất kiến trúc CNN "
        "phân cấp 3 mức (peak → locus → profile) với nhiều output phụ để "
        "giải thích. Repo đã clone kiến trúc đó (nhánh `simple` và `full`). "
        "Phiên bản tiếp theo trong repo là NoCFormer — thay backbone CNN "
        "bằng hierarchical set-Transformer + CORN ordinal head + synthetic "
        "mixture augmentation đơn giản.")

    # ==================== 3. CHẨN ĐOÁN ====================
    add_heading(doc, "3. Chẩn đoán: vì sao baseline chưa đủ tốt", level=1)
    add_para(doc,
        "Phân tích `results/history_nocformer.json` và metrics các model cũ:")
    add_table(doc,
        header=["Model", "Test Acc", "Split", "Ghi chú"],
        rows=[
            ["MAC (rule-based)",  "0.560",  "alternating",  "leak — không honest"],
            ["Random Forest",     "0.761",  "alternating",  "leak"],
            ["deepNoC simple",    "0.253",  "alternating",  "leak; chưa tune"],
            ["deepNoC full",      "0.820",  "alternating",  "leak — số bị thổi"],
            ["NoCFormer",         "0.436",  "grouped",      "honest, overfit nặng"],
        ])
    add_para(doc,
        "Ba vấn đề lớn được xác định:")

    add_para(doc, "(a) Split không nhất quán", bold=True)
    add_para(doc,
        "PROVEDIt tái-inject cùng một mẫu sinh học ở nhiều thời điểm và lần "
        "khuếch đại khác nhau. Split `alternating` đặt các replicate kề nhau "
        "vào train và test → leak. Tất cả số trên 0.56 trong các model cũ "
        "đều bị thổi phồng. Chỉ NoCFormer chạy với split `grouped` (test=1863) "
        "là honest, nhưng số đó (0.436) cũng cho thấy model overfit nặng "
        "(train acc ~91%, test acc 39–45%).")

    add_para(doc, "(b) Dữ liệu multi-contributor thiếu trầm trọng", bold=True)
    add_para(doc,
        "Mỗi class NoC=2..5 chỉ có ~160-176 profile gốc. Sau khi split "
        "grouped (chia theo pedigree), số training samples thật của các class "
        "hiếm còn lại rất ít. Sampler cân bằng class chỉ là band-aid — "
        "vẫn dùng lại cùng một vài profile thật.")

    add_para(doc, "(c) Pooling và inductive bias chưa phù hợp", bold=True)
    add_para(doc,
        "NoCFormer pool peak-axis bằng CLS-token + softmax attention. Softmax "
        "chuẩn hóa tổng về 1, làm mất tín hiệu \"có bao nhiêu peak được "
        "highlight\" — vốn là tín hiệu trực tiếp nhất để đếm contributor. "
        "Ngoài ra cross-peak attention chưa có bias đặc trưng theo "
        "allele-distance, nên model không biết đâu là quan hệ stutter "
        "(±1, ±2, ±0.2 repeat) cho đến khi tự học từ dữ liệu.")

    # ==================== 4. KIẾN TRÚC NoCNet-v2 ====================
    add_heading(doc, "4. Kiến trúc NoCNet-v2", level=1)
    add_para(doc,
        "NoCNet-v2 sửa đúng ba vấn đề chẩn đoán ở mục 3, đồng thời giữ "
        "đầu vào `[24 × 50 × 89]` tương thích pipeline dữ liệu hiện tại.")
    add_para(doc, "Sơ đồ tổng thể:")
    add_code(doc,
        "[B, 24, 50, 89] peaks\n"
        "  -> PeakEmbedder (89 -> 96)\n"
        "  -> 2x StutterBiasAttention   (attn logits + MLP(allele_i - allele_j))\n"
        "  -> DeepSetsPool [sum, max, log1p(count)]    <- count-preserving\n"
        "  -> 2x CrossLocusTransformer   (dye + locus pos)\n"
        "  -> ProfilePool [sum, mean, max]\n"
        "  -> CountAwareHead -> {softmax, scalar, CORN}   <- ensemble 3 view\n"
        "  + aux per-locus n_alleles head (chỉ học khi sample là synthetic)")
    add_para(doc, "Mô hình ~574k params, peak VRAM ~1.3 GiB ở batch 32 — vừa "
                  "GPU 6 GB. Chi tiết từng khối:")

    add_heading(doc, "4.1 PeakEmbedder", level=2)
    add_para(doc,
        "MLP 2 lớp 89 → 96 → 96 với GELU + LayerNorm + Dropout. "
        "Mục đích: chuyển vector feature 89-d sang không gian embedding "
        "đồng nhất cho các block sau.")

    add_heading(doc, "4.2 StutterBiasAttention", level=2)
    add_para(doc,
        "Multi-head self-attention thông thường, nhưng logits attention được "
        "cộng thêm một bias matrix:")
    add_code(doc,
        "bias_{i,j,h} = MLP_h(allele_j - allele_i)\n"
        "logits = Q K^T / sqrt(d_h) + bias")
    add_para(doc,
        "MLP 2 lớp 1 → 16 → n_heads, đầu vào là Δ allele giữa hai peak. "
        "Hệ quả: model tự học một \"prior\" cho thấy Δ = -1, -2, +1, ±0.2 "
        "là quan hệ stutter quan trọng — chính những giá trị mà data_loader "
        "đang encode bằng tay vào feature 29-76. Khi prior được học, model "
        "có thể bám vào quan hệ peak ngay cả khi feature 29-76 bị nhiễu "
        "(thay đổi sau khi sinh stutter synthetic).")
    add_para(doc,
        "Padding peak được mask qua key_padding_mask theo convention PyTorch "
        "(True = ignore). Hàng all-pad được clamp NaN về 0 để tránh lan "
        "instability.")

    add_heading(doc, "4.3 DeepSetsPool: count-preserving pooling", level=2)
    add_para(doc,
        "Đây là thay đổi kiến trúc QUAN TRỌNG NHẤT. Sau khi peak-axis "
        "self-attention xong, NoCNet-v2 không dùng CLS token mà dùng "
        "tổng hợp theo Deep Sets (Zaheer et al., 2017):")
    add_code(doc,
        "phi: per-peak MLP\n"
        "h_i' = phi(h_i)\n"
        "agg = concat( sum_i(h_i' * m_i),\n"
        "              max_i(h_i'),\n"
        "              log1p(count) / log(P+1) )\n"
        "locus_token = rho(agg)")
    add_para(doc,
        "Cả `sum` và `count` được giữ rõ ràng. Trong khi CLS-token attention "
        "tính softmax(QK^T), chuẩn hóa tổng attention weights về 1, làm "
        "mờ thông tin \"có 4 peak hay 8 peak\". DeepSetsPool đưa thông tin "
        "đếm trở thành một feature trực tiếp cho head phía sau.")

    add_heading(doc, "4.4 CrossLocusTransformer", level=2)
    add_para(doc,
        "Transformer 2 layer trên 24 token locus, có dye embedding + locus "
        "positional embedding. Mục đích: model biết mỗi locus thuộc kênh "
        "huỳnh quang nào và đứng ở vị trí nào trong kit, từ đó học mối "
        "quan hệ giữa các locus (ví dụ: locus có dropout cao thường có "
        "size bp lớn).")

    add_heading(doc, "4.5 ProfilePool + CountAwareHead", level=2)
    add_para(doc,
        "Profile pool: concat `[sum, mean, max]` qua 24 token locus, sau "
        "đó MLP về `d_model`. Head dùng 3 view song song:")
    add_bullet(doc, "Softmax classifier (CE + label-smoothing 0.05)")
    add_bullet(doc, "Scalar regression (smooth-L1) cho phép phạt off-by-one nhẹ")
    add_bullet(doc, "CORN ordinal (rank-monotonic) — vẫn từ paper "
                    "Cao, Mirjalili, Raschka (2020), nhưng KHÔNG stack focal "
                    "loss để tránh instability quan sát thấy ở NoCFormer.")
    add_para(doc,
        "Inference: ba probability vector từ 3 head được trung bình "
        "(`(p_cls + p_corn + p_reg) / 3`). Scalar regression được chuyển "
        "về phân phối class qua Gaussian kernel quanh giá trị dự đoán.")
    add_image(doc, os.path.join(here, "nocformer_arch.png"),
              width_in=6.2,
              caption="Hình 2. Sơ đồ pipeline (giữ cấu trúc hierarchical từ "
                      "NoCFormer; NoCNet-v2 thay PeakSetEncoder bằng "
                      "StutterBiasAttention + DeepSetsPool và thêm "
                      "scalar + ensemble head).")

    # ==================== 5. DỮ LIỆU TỔNG HỢP ====================
    add_heading(doc, "5. Dữ liệu tổng hợp: physics-based synthetic mixtures",
                level=1)
    add_para(doc,
        "Nguồn nghẹt cổ chai lớn nhất là 2-5P chỉ ~160-176 profile gốc. "
        "NoCNet-v2 tận dụng pool 2712 profile NoC=1 của PROVEDIt: một "
        "electropherogram k-contributor là (ở mức xấp xỉ bậc 1) tổng "
        "superposition của k profile single-source theo trọng số khối lượng "
        "DNA. Module `src/synth.py` thực hiện chính xác việc này, có physics "
        "đầy đủ.")

    add_heading(doc, "5.1 Superposition + Dirichlet weights", level=2)
    add_para(doc,
        "Với mỗi mẫu synthetic NoC=k:")
    add_bullet(doc, "Chọn k profile NoC=1 ngẫu nhiên (không trùng).")
    add_bullet(doc, "Lấy weight w ~ Dirichlet(α=1.5) — α=1.5 cho hỗn hợp "
                    "không quá lệch (α<1 thường tạo mixture rất lệch, "
                    "α>>1 tạo mixture cân bằng).")
    add_bullet(doc, "Với mỗi locus, gom (allele, height) từ tất cả contributor "
                    "rồi cộng dồn height theo trọng số.")

    add_heading(doc, "5.2 Tái sinh artefact (stutter / dropout / noise)", level=2)
    add_para(doc,
        "Sau bước superposition, ta thêm artefact thực tế cho từng peak "
        "thật:")
    add_table(doc,
        header=["Artefact", "Vị trí", "Tỉ lệ kỳ vọng", "CV"],
        rows=[
            ["Back stutter",      "allele - 1",   "0.10",   "0.20"],
            ["Double-back",       "allele - 2",   "0.01",   "0.30"],
            ["Forward stutter",   "allele + 1",   "0.03",   "0.30"],
            ["0.2 stutter",       "allele - 0.2", "0.005",  "0.40"],
        ])
    add_para(doc,
        "Allelic dropout: mỗi peak có xác suất rớt phụ thuộc chiều cao:")
    add_code(doc,
        "p_drop(h) = 0.30 * exp(-h / 250)")
    add_para(doc,
        "Drop-in noise: `Poisson(λ=0.5)` peak nhiễu mỗi locus, allele ngẫu "
        "nhiên trong dải kit, height đều trong [30, 80] RFU.")
    add_para(doc,
        "Stutter của contributor A có thể trùng allele thật của contributor B → "
        "rebucket tất cả peak ở cùng allele (gộp height) trước khi lọc theo "
        "LOD (50 RFU mặc định).")
    add_para(doc,
        "Quan trọng: nhãn ground-truth `n_alleles per locus` được snapshot "
        "TRƯỚC khi thêm artefact, để head phụ học dự đoán \"có bao nhiêu "
        "allele thật\" mà bỏ qua nhiễu — đúng signal model cần.")

    add_heading(doc, "5.3 Output", level=2)
    add_code(doc,
        "data/synthetic/X.npy            [N, 24, 50, 89] float32\n"
        "data/synthetic/y.npy            [N]              int64   (1..max_noc)\n"
        "data/synthetic/mix.npy          [N, 10]          float32 (sorted desc)\n"
        "data/synthetic/locus_nall.npy   [N, 24]          int8    (n_alleles)")
    add_para(doc,
        "Đọc bằng `np.memmap` → không tốn RAM. Mặc định `--n 30000`, "
        "lên 16 GB box có thể `--n 100000` thậm chí 200000.")

    # ==================== 6. LOSS ====================
    add_heading(doc, "6. Loss đa nhiệm", level=1)
    add_code(doc,
        "total = w_cls  * CE_label_smooth(class, NoC, class_weights)\n"
        "      + w_reg  * SmoothL1(scalar, NoC)\n"
        "      + w_corn * CORN_BCE(corn_logits, NoC)\n"
        "      + w_mix  * KL(mix_pred || mix_true)        [synthetic only]\n"
        "      + w_nall * CE(locus_nall_pred, n_alleles)  [synthetic only]")
    add_table(doc,
        header=["Tên", "Trọng số", "Đối tượng"],
        rows=[
            ["cls",   "1.0",  "Cross-entropy phân loại NoC, có "
                              "class-balanced weights (Cui 2019, β=0.999)"],
            ["reg",   "0.3",  "Smooth-L1 cho scalar regression"],
            ["corn",  "0.5",  "CORN BCE cumulative — KHÔNG focal stacking"],
            ["mix",   "0.1",  "KL chỉ tính trên row synthetic"],
            ["nall",  "0.1",  "CE per-locus chỉ trên row synthetic"],
        ])
    add_para(doc,
        "Điểm khác biệt vs NoCFormer: bỏ focal loss kết hợp class-balanced. "
        "Trong NoCFormer, hai cơ chế reweight chồng nhau gây gradient không "
        "ổn định ở các class hiếm; trong NoCNet-v2 chỉ giữ class-balanced "
        "+ label smoothing.")

    # ==================== 7. PIPELINE TRAIN ====================
    add_heading(doc, "7. Pipeline huấn luyện", level=1)
    add_heading(doc, "7.1 Hybrid loader: synthetic + real", level=2)
    add_para(doc,
        "Mỗi batch sample từ hai nguồn theo tỉ lệ `--p-synth` (mặc định 0.8):")
    add_bullet(doc, "Real PROVEDIt: lấy mẫu theo class-balanced weights "
                    "(class hiếm được lấy nhiều hơn).")
    add_bullet(doc, "Synthetic: lấy mẫu đều từ pool memmap.")
    add_para(doc,
        "Augmentation per-sample: log-normal height jitter (σ=0.10), "
        "random peak dropout (p=0.03), peak-axis shuffle. Tất cả vector hoá "
        "trên CPU per-batch, không tốn GPU.")

    add_heading(doc, "7.2 Optimizer + scheduler", level=2)
    add_bullet(doc, "AdamW, lr=3e-4, weight_decay=5e-4")
    add_bullet(doc, "Cosine LR schedule với linear warmup `warmup_epochs=3`")
    add_bullet(doc, "Gradient clipping max-norm 1.0")
    add_bullet(doc, "Early stopping patience 12 epoch")

    add_heading(doc, "7.3 Stochastic Weight Averaging", level=2)
    add_para(doc,
        "SWA (Izmailov et al., 2018) tự kích hoạt trong 25% epoch cuối "
        "(`swa_frac=0.25`). Mỗi epoch trong giai đoạn này, "
        "`AveragedModel.update_parameters(model)` cập nhật running average. "
        "Cả model live và SWA đều được eval, và best checkpoint = "
        "max(live_acc, swa_acc). Flag `was_swa` được ghi vào file ckpt.")
    add_para(doc,
        "Lý do: SWA gần như miễn phí, cải thiện generalization khoảng 1-2% "
        "trên hầu hết model nhỏ-vừa, và tránh được nhược điểm \"pick a noisy "
        "epoch\" của early stopping thuần.")

    add_heading(doc, "7.4 Fine-tune real-data", level=2)
    add_para(doc,
        "Sau khi pretrain xong, lệnh `python main.py finetune` mở checkpoint, "
        "train tiếp NHƯNG chỉ trên real PROVEDIt (`p_synth=0`), low LR "
        "(1e-5 mặc định), aug nhẹ hơn. Một bước an toàn quan trọng: "
        "trước khi train, đo baseline accuracy của checkpoint pretrain ngay, "
        "và best_ckpt khởi tạo từ baseline đó — nghĩa là fine-tune không "
        "bao giờ làm cho ckpt cuối cùng tệ hơn ckpt pretrain.")
    add_para(doc,
        "Tuỳ chọn `--freeze-peak` đóng băng peak embedder + peak attention, "
        "chỉ tune cross-locus + heads. Hữu ích khi real data quá ít, để "
        "tránh phá biểu diễn peak đã học từ synthetic.")

    # ==================== 8. ĐÁNH GIÁ ====================
    add_heading(doc, "8. Đánh giá: 5-fold grouped cross-validation", level=1)
    add_para(doc,
        "Module `src/cv.py` chạy 5-fold grouped CV trên cùng split cho "
        "nhiều model. Group key được trích từ pedigree trong sample name "
        "PROVEDIt (`src/split.py:_pedigree_key`) — đảm bảo các replicate "
        "cùng mẫu sinh học không nằm cả ở train và test.")
    add_para(doc,
        "Metrics báo cáo cho từng fold + summary (mean ± std):")
    add_bullet(doc, "accuracy (micro)")
    add_bullet(doc, "MAE (mean absolute error)")
    add_bullet(doc, "off-by-one accuracy (|y - ŷ| ≤ 1)")
    add_bullet(doc, "macro-F1")
    add_bullet(doc, "per-class precision/recall/F1, confusion matrix")
    add_para(doc,
        "Lý do cần off-by-one: trong forensic, dự đoán NoC=3 cho mẫu thật "
        "NoC=4 là sai về số nhưng vẫn cho phép phần lớn downstream "
        "deconvolution chạy được. Off-by-one là proxy gần với cost thực tế.")
    add_para(doc, "Lệnh chạy:")
    add_code(doc,
        "python main.py cv --models mac rf deepnoc_full nocformer nocnet_v2 \\\n"
        "                  --folds 5 --epochs 80 --batch-size 16")

    # ==================== 9. KẾT LUẬN ====================
    add_heading(doc, "9. Kết luận và roadmap", level=1)
    add_para(doc,
        "NoCNet-v2 là bước cải tiến có chủ ý: mỗi thay đổi (DeepSets, "
        "stutter-bias, 3-head ensemble, physics synth, SWA, fine-tune, "
        "grouped CV) đều khắc phục một vấn đề được chẩn đoán cụ thể trên "
        "code và metrics hiện có của repo.")
    add_para(doc, "Trạng thái:", bold=True)
    add_bullet(doc, "Architecture + losses + training loop + CV runner + "
                    "synthetic generator đã code xong, smoke-tested.")
    add_bullet(doc, "Smoke test trên fake data + GPU: forward / backward / "
                    "SWA / fine-tune đều hoạt động.")
    add_bullet(doc, "Chưa chạy full train trên GPU 16 GB (đợi máy ở nhà).")
    add_para(doc, "Roadmap đề xuất:", bold=True)
    add_bullet(doc, "Phase 1 — chạy `cv` honest với split grouped để có "
                    "mốc thật của baseline + model cũ.")
    add_bullet(doc, "Phase 2 — `synth --n 30000`, `train --model nocnet_v2 "
                    "--epochs 100`, `finetune --epochs 40`.")
    add_bullet(doc, "Phase 3 — nếu chưa đạt 0.85: tăng synth lên 100000, "
                    "ensemble nhiều seed, tune threshold per-class, "
                    "pseudo-labeling profile high-confidence.")
    add_para(doc,
        "Kỳ vọng hợp lý sau Phase 2: 80-85% accuracy honest trên grouped "
        "CV, off-by-one acc 95+%, MAE ~0.2-0.3. Mục tiêu 0.9 cần Phase 3 "
        "đầy đủ và may mắn với dữ liệu PROVEDIt không quá noisy nhãn.")

    # ==================== APPENDIX ====================
    add_heading(doc, "Phụ lục A. Tham số mặc định", level=1)
    add_table(doc,
        header=["Tham số", "Default", "Ý nghĩa"],
        rows=[
            ["d_model",            "96",      "Embedding dimension"],
            ["n_heads",            "4",       "Attention heads"],
            ["peak_layers",        "2",       "Layers nội locus"],
            ["locus_layers",       "2",       "Layers cross-locus"],
            ["dropout",            "0.15",    "Toàn bộ dropout"],
            ["batch_size",         "16",      "Default 6 GB VRAM"],
            ["epochs",             "80",      "Train, có early-stop"],
            ["lr",                 "3e-4",    "AdamW LR"],
            ["weight_decay",       "5e-4",    "AdamW WD"],
            ["warmup_epochs",      "3",       "Linear LR warmup"],
            ["swa_frac",           "0.25",    "Tail epoch dùng SWA"],
            ["p_synth",            "0.8",     "% batch là synthetic"],
            ["samples_per_epoch",  "4000",    "Hybrid loader iter/epoch"],
            ["jitter_sigma",       "0.10",    "Log-normal height jitter"],
            ["dropout_p",          "0.03",    "Random peak dropout"],
        ])

    add_heading(doc, "Phụ lục B. Lệnh chạy đầy đủ", level=1)
    add_code(doc,
        "# 1. Prepare data\n"
        "python main.py prepare \\\n"
        "    --data-dir 'data/provedit_processed/PROVEDIt_1-5-Person CSVs Filtered'\n"
        "\n"
        "# 2. Synthetic pool (1 lần, 5-15 phút)\n"
        "python main.py synth --n 30000 --max-noc 5\n"
        "\n"
        "# 3. Pretrain NoCNet-v2\n"
        "python main.py train --model nocnet_v2 --epochs 100 --batch-size 16 \\\n"
        "    --p-synth 0.85 --samples-per-epoch 4000 --split grouped\n"
        "\n"
        "# 4. Fine-tune trên real\n"
        "python main.py finetune --checkpoint results/best_nocnet_v2.pt \\\n"
        "    --epochs 40 --lr 1e-5 --tag nocnet_v2_ft\n"
        "\n"
        "# 5. 5-fold grouped CV apples-to-apples\n"
        "python main.py cv --models mac rf deepnoc_full nocformer nocnet_v2 \\\n"
        "    --folds 5 --epochs 80")

    add_heading(doc, "Phụ lục C. Tham chiếu", level=1)
    add_bullet(doc, "Taylor, J. & Humphries, M. (2024). \"deepNoC: A deep "
                    "learning system for the prediction of NoC from STR "
                    "DNA profiles.\" arXiv:2412.09803")
    add_bullet(doc, "Zaheer, M. et al. (2017). \"Deep Sets.\" NeurIPS.")
    add_bullet(doc, "Cao, W., Mirjalili, V., & Raschka, S. (2020). \"Rank "
                    "Consistent Ordinal Regression for Neural Networks "
                    "with Application to Age Estimation.\" Pattern "
                    "Recognition Letters.")
    add_bullet(doc, "Izmailov, P. et al. (2018). \"Averaging Weights Leads "
                    "to Wider Optima and Better Generalization\" "
                    "(Stochastic Weight Averaging). UAI.")
    add_bullet(doc, "Cui, Y. et al. (2019). \"Class-Balanced Loss Based on "
                    "Effective Number of Samples.\" CVPR.")
    add_bullet(doc, "PROVEDIt dataset: https://lftdi.camden.rutgers.edu/provedit/")

    out_path = os.path.join(here, "NoCNet_v2_report.docx")
    doc.save(out_path)
    print(f"[ok] wrote {out_path}")


if __name__ == "__main__":
    build()
