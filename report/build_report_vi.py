"""
Sinh báo cáo NoCFormer bằng tiếng Việt (.docx) — phiên bản đầy đủ:
  - giải thích tập dữ liệu cho người mới (PROVEDIt, GeneMapper, locus, dye, NoC),
  - nhúng ảnh kiến trúc + ảnh layout dataset,
  - giải thích từng module trong NoCFormer (PeakEmbedder, PeakSetEncoder,
    PeakHead, LocusTransformer, LocusHead, ProfileHead, CORN head),
  - trình bày pipeline cải tiến (augment, loss, MC-Dropout TTA, split).

Chạy:
    python report/build_report_vi.py
"""

from __future__ import annotations

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# -------------------- Style helpers --------------------------------------

def set_default_font(doc: Document, name: str = "Calibri", size_pt: int = 11):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)
    rpr = style.element.rPr
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x2D, 0x5C)
    return h


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)


def add_code(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    rpr = r._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")
    rpr.append(rfonts)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)


def add_table(doc: Document, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val
    return table


def add_image(doc: Document, path: str, width_in: float = 6.5, caption: str = ""):
    if not os.path.exists(path):
        add_para(doc, f"[Không tìm thấy ảnh: {path}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)


# -------------------- Build report ---------------------------------------

def build():
    here = os.path.dirname(__file__)
    doc = Document()
    set_default_font(doc)

    # ==================== TRANG TIÊU ĐỀ ====================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("NoCFormer")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1F, 0x2D, 0x5C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Mô hình cải tiến cho bài toán xác định "
                    "số người đóng góp DNA (NoC) từ hồ sơ STR")
    r.italic = True
    r.font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Báo cáo cải tiến cho deepNoC — Taylor & Humphries (2024), "
                 "arXiv:2412.09803\n").italic = True
    meta.add_run("Repo: deepNoC (mã nguồn hiện tại)\n")
    meta.add_run("Trạng thái: thiết kế + code đã xong; "
                 "đợi train full trên GPU\n")

    doc.add_paragraph()

    # ==================== 1. TÓM TẮT ====================
    add_heading(doc, "1. Tóm tắt", level=1)
    add_para(doc,
        "Báo cáo này đề xuất NoCFormer — mô hình deep learning mới để dự đoán "
        "số người đóng góp DNA (NoC, Number of Contributors) cho hồ sơ STR. "
        "Nền tảng so sánh là deepNoC của Taylor & Humphries (2024); repo hiện "
        "tại đã clone cấu trúc đó. NoCFormer giữ lại phần đầu vào "
        "[24 × 50 × 89] và triết lý multi-output để giải thích, nhưng thay "
        "backbone CNN bằng hierarchical set-Transformer, thay softmax phân "
        "loại bằng đầu ra ordinal CORN, và thay bộ giả lập GAN 100k profile "
        "bằng pipeline tổng hợp mixture trực tiếp từ profile single-source "
        "trên PROVEDIt.")
    add_para(doc, "Sáu cải tiến chính:", bold=True)
    add_bullet(doc, "(1) Bộ encode peak có tính chất hoán vị bất biến — "
                    "peak ở mỗi locus là một TẬP HỢP, không phải chuỗi.")
    add_bullet(doc, "(2) Self-attention liên locus có thêm dye-channel "
                    "embedding — nắm được pull-up và cân bằng dye.")
    add_bullet(doc, "(3) Đầu ra ordinal CORN — phạt nhẹ khi đoán nhầm gần, "
                    "phạt nặng khi đoán nhầm xa.")
    add_bullet(doc, "(4) Synthetic-mixture augmentation — sinh hồ sơ k người "
                    "từ k hồ sơ 1 người, không cần GAN.")
    add_bullet(doc, "(5) MC-Dropout test-time augmentation — phân phối "
                    "xác suất đã calibrate + entropy làm điểm tin cậy.")
    add_bullet(doc, "(6) Grouped-stratified split — chống rò rỉ giữa các "
                    "lần re-inject cùng một mixture sinh học.")

    # ==================== 2. BỐI CẢNH ====================
    add_heading(doc, "2. Bối cảnh bài toán", level=1)

    add_heading(doc, "2.1 STR DNA là gì?", level=2)
    add_para(doc,
        "STR (Short Tandem Repeat) là các đoạn DNA có lõi 2–6 nucleotide "
        "lặp lại nhiều lần liên tiếp. Số lần lặp ở mỗi vị trí (locus) khác "
        "nhau giữa các cá nhân, vì vậy STR là dấu vân tay di truyền chuẩn "
        "trong điều tra hình sự. Một locus thường có 2 allele (một từ bố, "
        "một từ mẹ), được biểu thị bằng số lần lặp.")

    add_heading(doc, "2.2 Hồ sơ STR đọc bằng máy như thế nào?", level=2)
    add_para(doc,
        "Mẫu DNA → tách → khuếch đại bằng PCR (kit GlobalFiler, Identifiler…) "
        "→ điện di mao quản (capillary electrophoresis) → mỗi đoạn STR "
        "phát huỳnh quang → phổ đo gọi là EPG (electropherogram). EPG là "
        "tập các đỉnh peak trên 5 dye color khác nhau, mỗi peak ứng với "
        "một allele tại một locus.")

    add_heading(doc, "2.3 NoC và vì sao nó quan trọng", level=2)
    add_para(doc,
        "NoC (Number of Contributors) = số người đóng góp DNA vào mẫu. Đây "
        "là bước ĐẦU TIÊN trong mọi quy trình giải mã hồ sơ trộn (mixture). "
        "Nếu đếm thiếu NoC → người thật bị loại oan; nếu đếm thừa → người "
        "không liên quan bị nghi sai. Ở mức ≥ 3 người, ngay cả chuyên gia "
        "cũng đoán sai vì các allele chồng chéo. Đó là lý do cần một mô "
        "hình học sâu xác định NoC.")

    # ==================== 3. DATASET ====================
    add_heading(doc, "3. Dataset PROVEDIt — giải thích chi tiết cho người mới",
                level=1)

    add_heading(doc, "3.1 PROVEDIt là gì?", level=2)
    add_para(doc,
        "PROVEDIt = Project Research Openness for Validation with Empirical "
        "Data. Đây là tập dữ liệu công khai do nhóm Catherine Grgicak "
        "(Rutgers University) phát hành năm 2017. Họ chuẩn bị các mẫu DNA "
        "từ 1, 2, 3, 4 và 5 người đóng góp với nhiều tỷ lệ khác nhau, chạy "
        "qua nhiều kit + nhiều máy + nhiều thời gian inject, rồi công bố "
        "toàn bộ EPG và file CSV phân tích. Đây là tập benchmark phổ biến "
        "nhất cho mọi thuật toán NoC.")

    add_heading(doc, "3.2 File CSV GeneMapper — đọc cột nào?", level=2)
    add_para(doc,
        "GeneMapper là phần mềm thương mại của ThermoFisher, đọc EPG thô "
        "rồi xuất ra bảng các peak. Bảng dạng wide (mỗi dòng là một locus "
        "trong một sample, có nhiều cột Allele 1..N, Size 1..N, Height 1..N). "
        "Loader của repo đọc cả CSV/XLSX:")
    add_bullet(doc, "Sample Name / SampleName / Sample File / Sample → tên sample")
    add_bullet(doc, "Marker / Locus → tên locus")
    add_bullet(doc, "Dye / Color → kênh màu (5 kênh: Blue, Green, Yellow, Red, Purple)")
    add_bullet(doc, "Allele 1, Allele 2, … → giá trị allele (số repeat, có thể là 'X', 'Y' với AMEL)")
    add_bullet(doc, "Size 1, Size 2, … → kích thước fragment tính bằng base pair (bp)")
    add_bullet(doc, "Height 1, Height 2, … → chiều cao peak tính bằng RFU (Relative Fluorescence Unit)")
    add_para(doc,
        "Loader chuyển wide → long format: mỗi peak một dòng với các cột "
        "[SampleName, Marker, Dye, Allele, Size, Height]. Sau đó gom theo "
        "sample để build tensor.")

    add_heading(doc, "3.3 Quy tắc lọc peak", level=2)
    add_para(doc, "Loader bỏ qua các peak trong các trường hợp sau:")
    add_bullet(doc, "Allele rỗng / NaN (cột để trống).")
    add_bullet(doc, "Allele = ‘OL’ (off-ladder: peak không khớp ladder chuẩn → coi là noise).")
    add_bullet(doc, "Allele không parse được thành số (trừ ‘X’/‘Y’ ở locus AMEL).")
    add_bullet(doc, "Height ≤ 0.")
    add_bullet(doc, "Locus không thuộc 24 locus của GlobalFiler.")
    add_para(doc,
        "Riêng AMEL (Amelogenin) là locus xác định giới tính: ‘X’ → 1.0, "
        "‘Y’ → 2.0. Nam thường có cả X và Y, nữ chỉ có X.")

    add_heading(doc, "3.4 24 locus GlobalFiler (theo thứ tự)", level=2)
    add_table(doc,
        ["Stt", "Locus", "Stt", "Locus"],
        [
            ["1", "D3S1358", "13", "D19S433"],
            ["2", "vWA",     "14", "TH01"],
            ["3", "D16S539", "15", "FGA"],
            ["4", "CSF1PO",  "16", "D22S1045"],
            ["5", "TPOX",    "17", "D5S818"],
            ["6", "Y-Indel", "18", "D13S317"],
            ["7", "AMEL",    "19", "D7S820"],
            ["8", "D8S1179", "20", "SE33"],
            ["9", "D21S11",  "21", "D10S1248"],
            ["10", "D18S51", "22", "D1S1656"],
            ["11", "DYS391", "23", "D12S391"],
            ["12", "D2S441", "24", "D2S1338"],
        ])
    add_para(doc,
        "‘D’ + số nhiễm sắc thể + ‘S’ + ID locus là quy ước đặt tên: "
        "ví dụ D3S1358 = locus 1358 trên nhiễm sắc thể 3. Các locus có "
        "tên riêng (vWA, FGA, TPOX…) là các marker truyền thống có lịch "
        "sử lâu hơn. Y-Indel và DYS391 nằm trên nhiễm sắc thể Y → chỉ "
        "phát hiện ở nam. Thứ tự trên rất quan trọng vì nó được dùng "
        "trực tiếp làm one-hot encoding (24 cột đầu trong 89 features).")

    add_heading(doc, "3.5 5 dye channel (kênh màu)", level=2)
    add_para(doc,
        "Mỗi locus được gán cho một dye color sao cho khi điện di, các "
        "fragment cùng kích thước (cùng bp) nhưng khác dye sẽ phân biệt "
        "được bằng máy quang phổ:")
    add_table(doc,
        ["Dye", "Màu", "Locus thuộc dye"],
        [
            ["B", "Blue",   "D3S1358, vWA, D16S539, CSF1PO, TPOX"],
            ["G", "Green",  "Y-Indel, AMEL, D8S1179, D21S11, D18S51, DYS391"],
            ["Y", "Yellow", "D2S441, D19S433, TH01, FGA"],
            ["R", "Red",    "D22S1045, D5S818, D13S317, D7S820, SE33"],
            ["P", "Purple", "D10S1248, D1S1656, D12S391, D2S1338"],
        ])
    add_para(doc,
        "Pull-up artefact (huỳnh quang lan từ dye này sang dye khác) "
        "thường xảy ra trong cùng một dye → đây là lý do NoCFormer dùng "
        "‘dye embedding’ trong LocusTransformer.")

    add_heading(doc, "3.6 Suy ra NoC từ tên file / sample", level=2)
    add_para(doc,
        "PROVEDIt mã hóa pedigree (ai, tỷ lệ ra sao) ngay trong tên file "
        "và tên sample. Loader trong repo có 2 lớp suy luận:")
    add_para(doc, "(a) Từ tên file:", bold=True)
    add_bullet(doc, "‘_1P.csv’ → NoC=1; ‘_2P.csv’ → NoC=2; … ‘_5P.csv’ → NoC=5")
    add_bullet(doc, "‘_2-5P.csv’ → file chứa hỗn hợp 2..5 người, phải parse "
                    "tiếp ở mức sample.")
    add_para(doc, "(b) Từ tên sample (cho file multi-NoC):", bold=True)
    add_bullet(doc, "‘…-1;1-…’       → NoC = 2  (2 nhóm tỷ lệ → 2 người)")
    add_bullet(doc, "‘…-1;2;1-…’     → NoC = 3")
    add_bullet(doc, "‘…-1;1;1;1-…’   → NoC = 4")
    add_bullet(doc, "‘…-1;1;1;1;1-…’ → NoC = 5")
    add_bullet(doc, "Fallback cũ: ‘1to1’, ‘1to1to1’, hoặc ‘2p’/‘3p’/‘4p’/‘5p’")

    add_heading(doc, "3.7 Tensor đầu vào [24 × 50 × 89]", level=2)
    add_para(doc, "Mỗi profile được đóng gói thành một tensor 3 chiều:")
    add_bullet(doc, "Trục 0 = 24 locus của GlobalFiler (theo thứ tự cố định ở mục 3.4).")
    add_bullet(doc, "Trục 1 = tối đa 50 peak ở mỗi locus. Nếu ít hơn 50 → zero-pad. "
                    "Nếu nhiều hơn → giữ 50 peak đầu sau khi sắp xếp theo Size.")
    add_bullet(doc, "Trục 2 = 89 đặc trưng cho mỗi peak (chia 3 nhóm, xem mục 3.8).")
    add_image(doc, os.path.join(here, "dataset_layout.png"),
              width_in=6.5, caption="Hình 1 — Layout tensor đầu vào và "
              "decoder tên sample PROVEDIt")

    add_heading(doc, "3.8 89 đặc trưng — giải thích từng nhóm", level=2)
    add_para(doc, "Nhóm 1 — Định danh peak (idx 0..28):", bold=True)
    add_bullet(doc, "0..23 : one-hot 24 locus (chỉ 1 trong 24 ô = 1).")
    add_bullet(doc, "24    : allele / 100 (chuẩn hoá; ví dụ allele 15 → 0.15).")
    add_bullet(doc, "25    : size / 100 (kích thước fragment, bp).")
    add_bullet(doc, "26    : height / 33000 (chiều cao peak, RFU; 33000 ≈ saturation).")
    add_bullet(doc, "27    : allele frequency (tần số allele trong quần thể; "
                    "repo hiện dùng giá trị mặc định 0.01).")
    add_bullet(doc, "28    : peak label probability (xác suất peak là thật, "
                    "không phải artefact). Repo dùng heuristic vì chưa có MHCNN.")

    add_para(doc, "Nhóm 2 — Quan hệ stutter (idx 29..76):", bold=True)
    add_para(doc,
        "Stutter là peak phụ xuất hiện do trượt PCR — thường thấp hơn peak "
        "‘parent’ và ở vị trí allele lệch ±1 hoặc ±2 repeat. NoCFormer "
        "biết đến 4 loại stutter:")
    add_bullet(doc, "back stutter (-1 repeat) — hay gặp nhất, tỉ lệ ≈ 10%")
    add_bullet(doc, "double-back stutter (-2 repeat) — hiếm, tỉ lệ ≈ 1%")
    add_bullet(doc, "forward stutter (+1 repeat) — tỉ lệ ≈ 3%")
    add_bullet(doc, "0.2-repeat stutter (lệch 0.2) — tỉ lệ ≈ 0.5%")
    add_para(doc,
        "Mỗi loại được mô tả bằng 6 giá trị (parent allele, parent height, "
        "tỷ lệ height, tỷ lệ kỳ vọng, allele frequency của parent, plp của "
        "parent), và lưu cho 2 chiều: ‘peak này là stutter của peak khác’ + "
        "‘peak này là parent của một stutter’. Tổng = 4 × 6 × 2 = 48 ô "
        "(idx 29..76).")

    add_para(doc, "Nhóm 3 — Context (idx 77..88):", bold=True)
    add_bullet(doc, "77 : số peak ở locus / 100 (giống nhau cho mọi peak trong locus).")
    add_bullet(doc, "78 : số peak toàn profile / 1000.")
    add_bullet(doc, "79..88 : mixture proportions — tỷ lệ DNA của 10 người "
                    "đóng góp lớn nhất, tính bằng smart-start algorithm "
                    "(STRmix). Repo dùng phiên bản đơn giản hoá.")

    add_heading(doc, "3.9 Thống kê dataset hiện tại", level=2)
    add_para(doc, "Sau khi prepare với cấu hình max_1person=70 (nén bớt 1-người để giảm imbalance):")
    add_table(doc,
        ["NoC", "Số profile"],
        [
            ["1", "70"],
            ["2", "174"],
            ["3", "160"],
            ["4", "176"],
            ["5", "156"],
            ["Tổng", "736"],
        ])

    # ==================== 4. DEEPNOC GỐC + ĐIỂM YẾU ====================
    add_heading(doc, "4. deepNoC gốc và những điểm yếu", level=1)

    add_heading(doc, "4.1 Tóm tắt deepNoC", level=2)
    add_para(doc,
        "deepNoC dùng backbone CNN 16 lớp: per-peak conv1d → AdaptiveMaxPool "
        "→ per-locus conv1d → AdaptiveMaxPool → MLP cuối → softmax 1..10. "
        "Train trên 100,000 hồ sơ giả lập bằng simDNAmixtures + GAN, "
        "fine-tune trên 371 hồ sơ PROVEDIt. Test accuracy: 90% trên "
        "PROVEDIt 1-5 người, 72% trên giả lập 1-10 người.")

    add_heading(doc, "4.2 8 điểm yếu", level=2)
    add_table(doc,
        ["Vấn đề", "Tác động"],
        [
            ["Conv1D trên trục peak nhạy cảm với thứ tự peak",
             "Peak trong locus là TẬP HỢP, không có thứ tự tự nhiên. Repo "
             "đang sort theo Size → mô hình phải học để bỏ qua thứ tự."],
            ["Softmax phân loại 1..K bỏ qua tính ordinal",
             "Đoán NoC=4 thành NoC=9 bị phạt y hệt như đoán NoC=4 thành NoC=3, "
             "nên model không có động lực ‘đúng gần’."],
            ["AdaptiveMaxPool nén thông tin sớm",
             "Chi tiết tương tác giữa peak và đặc biệt là giữa các locus "
             "trong cùng một dye bị mất."],
            ["Không có tương tác cross-locus thực sự",
             "Pull-up artefact span dye; cân bằng dye thông tin tốt cho NoC. "
             "Conv1D kernel=3 chỉ thấy locus láng giềng theo INDEX, mà "
             "thứ tự index là tuỳ ý sinh học."],
            ["Mất cân bằng class nặng (≈80% là 1 người trên PROVEDIt gốc)",
             "Không reweight thì model dễ chỉ đoán NoC=1."],
            ["Split alternating (mỗi profile thứ 2)",
             "PROVEDIt re-inject cùng một mixture nhiều lần. Split alternating "
             "đặt các bản sao gần nhau lên 2 phía → leakage → accuracy "
             "test bị thổi phồng."],
            ["Một forward pass duy nhất khi inference",
             "Không có ước lượng độ tin cậy — trong khi paper gốc lập luận "
             "ngay trong Hình 7 rằng phải dùng probability threshold."],
            ["Phụ thuộc vào GAN simulator 100k profile",
             "Nếu không có simulator (như repo này), dữ liệu train chỉ "
             "vài trăm profile → quá nhỏ cho deep model."],
        ])

    # ==================== 5. KIẾN TRÚC NOCFORMER ====================
    add_heading(doc, "5. Kiến trúc NoCFormer", level=1)

    add_heading(doc, "5.1 Tổng thể", level=2)
    add_image(doc, os.path.join(here, "nocformer_arch.png"),
              width_in=6.5, caption="Hình 2 — Kiến trúc NoCFormer end-to-end")

    add_para(doc,
        "Pipeline gồm 6 module xếp tầng. Đầu vào tensor [B, 24, 50, 89]; đầu "
        "ra là 6 dictionary keys (1 chính = NoC, 5 phụ để giải thích). "
        "Mục con bên dưới giải thích từng module một.")

    # ---- 5.2 PeakEmbedder ----
    add_heading(doc, "5.2 PeakEmbedder — biến 89 feature thô thành embedding d-chiều",
                level=2)
    add_para(doc, "Vào: [B, 24, 50, 89]   Ra: [B, 24, 50, d=128]")
    add_para(doc,
        "Đây là MLP 2 lớp ẩn áp dụng độc lập lên TỪNG peak. Lý do dùng MLP "
        "thay vì Conv1D là vì 89 feature đã là vector kết hợp đầy đủ "
        "(one-hot + scalar + stutter), không phải chuỗi thời gian.")
    add_code(doc,
        "Sequential(\n"
        "  Linear(89 → 128), GELU, LayerNorm(128), Dropout(0.1),\n"
        "  Linear(128 → 128), GELU, LayerNorm(128)\n"
        ")")
    add_para(doc,
        "GELU mượt hơn ReLU, LayerNorm ổn định gradient cho Transformer ở "
        "downstream, Dropout giúp MC-Dropout hoạt động khi inference.")

    # ---- 5.3 PeakSetEncoder ----
    add_heading(doc,
        "5.3 PeakSetEncoder — Transformer per-locus với CLS token", level=2)
    add_para(doc, "Vào: [B, 24, 50, d]   Ra: locus_token [B, 24, d] + "
                  "peak_token [B, 24, 50, d]")
    add_para(doc,
        "Module này được gọi ‘song song’ trên 24 locus (reshape về "
        "[B*24, 50, d]). Mỗi locus có một CLS token học được; CLS này attend "
        "vào tất cả 50 peak và sẽ trở thành representation tổng hợp của "
        "locus.")
    add_para(doc, "Các bước:", bold=True)
    add_bullet(doc, "Tạo CLS token (parameter [1,1,d] init bằng truncated normal).")
    add_bullet(doc, "Concat CLS vào trước 50 peak token → chuỗi 51 token.")
    add_bullet(doc, "Tạo key_padding_mask: peak có height=0 (padding) bị "
                    "mask để attention bỏ qua.")
    add_bullet(doc, "Chạy 2 lớp Pre-LN Transformer block (LayerNorm → MultiheadAttention "
                    "→ residual → LayerNorm → FFN(4×d) → residual).")
    add_bullet(doc, "LayerNorm cuối → tách CLS thành locus_token, "
                    "phần còn lại là peak_token đã refine.")
    add_para(doc,
        "Tại sao set, không phải sequence? Vì peak ở locus không có thứ tự "
        "vốn có. Nếu dùng CNN/RNN, model sẽ học một thứ tự giả tạo (Size-sort). "
        "Self-attention permutation-equivariant + CLS → permutation-invariant.")

    # ---- 5.4 PeakHead ----
    add_heading(doc, "5.4 PeakHead — đầu ra phụ ở mức peak", level=2)
    add_para(doc, "Vào: peak_token [B*24, 50, d]   Ra: 2 đầu ra")
    add_bullet(doc, "prop_allelic [B, 24, 50, 1] : tỷ lệ peak là allelic "
                    "(không phải stutter / artefact). Sigmoid trong [0,1].")
    add_bullet(doc, "n_alleles [B, 24, 50, 21] : phân loại số allele tại "
                    "vị trí peak (0..20).")
    add_para(doc,
        "Đây là đầu ra GIẢI THÍCH — analyst nhìn vào để biết model nghĩ "
        "peak nào là thật, peak nào là noise.")

    # ---- 5.5 LocusTransformer ----
    add_heading(doc, "5.5 LocusTransformer — Transformer cross-locus", level=2)
    add_para(doc, "Vào: locus_token [B, 24, d]   Ra: profile_token [B, d] + "
                  "refined locus_token [B, 24, d]")
    add_para(doc, "Quy trình:", bold=True)
    add_bullet(doc, "Cộng learnable position embedding [1, 24, d] vào 24 locus token.")
    add_bullet(doc, "Cộng dye embedding (5 dye, mỗi dye một vector). "
                    "Locus cùng dye sẽ chia sẻ một thành phần biểu diễn → "
                    "model học được pattern dye-balance / pull-up.")
    add_bullet(doc, "Concat một profile CLS token (parameter [1,1,d]) vào trước.")
    add_bullet(doc, "Tạo key_padding_mask theo locus_active = "
                    "(any peak height > 0).")
    add_bullet(doc, "Chạy 4 lớp Pre-LN Transformer block.")
    add_bullet(doc, "Tách profile_token (CLS) ra; phần còn lại là 24 locus "
                    "token đã refine.")
    add_para(doc,
        "Việc chạy 4 lớp ở đây sâu hơn 2 lớp peak-level vì cross-locus "
        "interaction (chỗ thông tin NoC sống) cần nhiều round attention hơn.")

    # ---- 5.6 LocusHead ----
    add_heading(doc, "5.6 LocusHead — đầu ra phụ ở mức locus", level=2)
    add_para(doc, "Vào: locus_token đã refine   Ra: 2 đầu ra")
    add_bullet(doc, "n_alleles [B, 24, 20] : phân loại số allele ở mỗi "
                    "locus (1..20).")
    add_bullet(doc, "mix_props [B, 24, 10] : tỷ lệ đóng góp của 10 donor "
                    "lớn nhất ở mỗi locus, softmax tổng = 1.")

    # ---- 5.7 ProfileHead + CORN ----
    add_heading(doc, "5.7 ProfileHead — CORN ordinal NoC head", level=2)
    add_para(doc, "Vào: profile_token [B, d]   Ra: NoC logits [B, K-1] + "
                  "mix_props [B, 10]")
    add_para(doc,
        "ĐÂY LÀ HEAD CHÍNH. Khác biệt lớn nhất so với deepNoC: dùng CORN "
        "(Cao, Mirjalili, Raschka 2020) thay cho softmax phân loại.")
    add_para(doc, "Cơ chế CORN:", bold=True)
    add_bullet(doc, "Thay vì 10 logits (1..10), output K-1 = 9 logits. Logit thứ k "
                    "biểu diễn xác suất P(y > k).")
    add_bullet(doc, "Sigmoid mỗi logit rồi nhân tích lũy: cum[k] = σ(z₁)·σ(z₂)·…·σ(z_k). "
                    "Tích lũy đảm bảo TÍNH ĐƠN ĐIỆU: P(y>k) ≤ P(y>k-1) (không bao giờ "
                    "có nghịch lý).")
    add_bullet(doc, "Predict: NoC = (cum > 0.5).sum() + 1.")
    add_bullet(doc, "Loss: BCE riêng cho từng task k, với CONDITIONAL TRAINING "
                    "(task k chỉ dùng sample có target > k-1).")
    add_para(doc, "Lợi ích thực tế:", bold=True)
    add_bullet(doc, "Đoán NoC=4 thành NoC=3 chỉ phạt 1 task; thành NoC=9 phạt 5 task.")
    add_bullet(doc, "Output có ý nghĩa thống kê (xác suất tích lũy), không cần softmax.")
    add_bullet(doc, "Dễ thêm class-balanced + focal weighting vì BCE độc lập.")

    # ---- 5.8 corn_logits_to_class_probs ----
    add_heading(doc, "5.8 Hàm corn_logits_to_class_probs — chuyển CORN sang phân phối lớp",
                level=2)
    add_para(doc,
        "Để vẫn lấy được phân phối xác suất P(y=k) cho mọi k (cần cho "
        "threshold analysis của Hình 7), ta tính:")
    add_code(doc,
        "P(y >= k)  = [1, cum[1], cum[2], ..., cum[K-1]]\n"
        "P(y >  k)  = [cum[1], cum[2], ..., cum[K-1], 0]\n"
        "P(y == k)  = P(y >= k) - P(y > k)")
    add_para(doc,
        "Vì cum đơn điệu giảm theo k, hiệu này luôn ≥ 0 và tổng theo k bằng 1.")

    # ==================== 6. PIPELINE ====================
    add_heading(doc, "6. Pipeline cải tiến", level=1)

    add_heading(doc, "6.1 Synthetic-mixture augmentation (cốt lõi)", level=2)
    add_para(doc,
        "Một EPG từ k người đóng góp ≈ TỔNG SỐ HỌC của k EPG đơn nguồn, "
        "với các hệ số tỷ lệ ngẫu nhiên. Hàm synthetic_mix() khai thác "
        "điều này:")
    add_bullet(doc, "Sample tỷ lệ w ~ Dirichlet(1,…,1) trên simplex k-chiều.")
    add_bullet(doc, "Lấy k profile single-source ngẫu nhiên từ pool train.")
    add_bullet(doc, "Cộng dồn height của các peak cùng (locus, allele), "
                    "có nhân trọng số w.")
    add_bullet(doc, "Recompute các feature global (tổng peak/locus, "
                    "tổng peak/profile, mix_props).")
    add_bullet(doc, "Trả về tensor mới + nhãn k.")
    add_para(doc,
        "Với 70 profile single-source ta có thể sinh C(70, 2) ≈ 2,400 mixture 2-người, "
        "C(70, 3) ≈ 54,000 mixture 3-người, … về cơ bản là vô hạn dữ liệu mà "
        "không cần GAN. Trong training loop, p_mix=0.5 nghĩa là một nửa "
        "batch là synthetic, một nửa là PROVEDIt thật.")

    add_heading(doc, "6.2 Augmentation peak-level", level=2)
    add_bullet(doc, "peak_height_jitter: nhân height với LN(0, σ=0.12) → "
                    "mô phỏng biến thiên giữa các lần inject.")
    add_bullet(doc, "random_peak_dropout(p=0.03): zero-out ngẫu nhiên 3% "
                    "peak thật → mô phỏng allelic drop-out.")
    add_bullet(doc, "shuffle_peak_axis: hoán vị ngẫu nhiên peak trong locus. "
                    "AN TOÀN vì encoder permutation-invariant; dạy model "
                    "không phụ thuộc vào Size-sort.")

    add_heading(doc, "6.3 Class-balanced focal CORN loss", level=2)
    add_para(doc,
        "Mất cân bằng PROVEDIt được xử lý 2 lớp: (1) WeightedRandomSampler "
        "lấy mỗi class với xác suất bằng nhau khi sample batch; (2) "
        "class-balanced weight (Cui 2019, β=0.999) nhân với BCE của CORN; "
        "(3) focal modulator (γ=1) ưu tiên các sample khó.")

    add_heading(doc, "6.4 MC-Dropout test-time augmentation", level=2)
    add_para(doc,
        "Hàm predict_with_tta() bật dropout khi inference, jitter + shuffle "
        "input, chạy n_samples=20 forward pass, average xác suất. Trả về "
        "(probs, entropy, preds). Entropy là confidence score: entropy thấp "
        "→ model chắc chắn → có thể assign; entropy cao → abstain. Trùng "
        "khớp với use case Hình 7 của paper gốc.")

    add_heading(doc, "6.5 Grouped stratified split — chống leakage", level=2)
    add_para(doc,
        "Tên sample PROVEDIt nhúng pedigree (ví dụ ‘…-31_32-1;1-…’ là cùng "
        "một mixture sinh học giữa donor 31 và 32 tỷ lệ 1:1). Mọi replicate "
        "/ injection của cùng một mixture chia sẻ key này. Module src/split.py "
        "extract key bằng regex và dùng GroupShuffleSplit để tất cả "
        "replicate cùng mixture chỉ ở một phía split. Stratify theo NoC "
        "đảm bảo mỗi class có ở cả train và test.")

    # ==================== 7. CODE ====================
    add_heading(doc, "7. File mã nguồn đã thêm", level=1)
    add_table(doc,
        ["File", "Vai trò"],
        [
            ["models/nocformer/architecture.py",
             "PeakEmbedder, PeakSetEncoder, LocusTransformer, "
             "PeakHead, LocusHead, ProfileHead, NoCFormer, "
             "corn_logits_to_class_probs."],
            ["models/nocformer/losses.py",
             "corn_loss (CORN ordinal), class_balanced_weights, "
             "NoCFormerLoss (multi-task wrapper)."],
            ["models/nocformer/augment.py",
             "synthetic_mix, peak_height_jitter, random_peak_dropout, "
             "shuffle_peak_axis, TrainingAugmenter."],
            ["models/nocformer/train.py",
             "train_nocformer (AdamW + cosine + warmup + balanced sampler), "
             "predict_with_tta, load_nocformer."],
            ["src/split.py",
             "stratified_split, grouped_stratified_split (pedigree-aware)."],
            ["main.py",
             "Thêm --model nocformer; --split {alternating,stratified,grouped}; "
             "các knob model/aug."],
            ["report/build_report_vi.py",
             "Sinh báo cáo này."],
            ["report/draw_architecture.py / draw_dataset.py",
             "Sinh các hình minh hoạ trong báo cáo."],
        ])

    # ==================== 8. HYPERPARAMS ====================
    add_heading(doc, "8. Hyperparameter mặc định", level=1)
    add_table(doc,
        ["Tham số", "Giá trị", "Lý do"],
        [
            ["d_model", "128", "Đủ cho 1.3M-param model trên 736 PROVEDIt + augmentation."],
            ["n_heads", "4", "Tỷ lệ chuẩn d_model / 32 cho Transformer nhỏ."],
            ["peak_layers", "2", "Self-attention 2 round trên 50 token là đủ."],
            ["locus_layers", "4", "Cross-locus interaction quan trọng hơn → sâu hơn."],
            ["dropout", "0.15", "Vừa regularise vừa MC-Dropout TTA."],
            ["epochs", "200", "Khớp paper gốc; cosine schedule + warmup 5 epoch."],
            ["batch_size", "32", "Vừa GPU 16 GB ở d_model=128."],
            ["lr", "3e-4", "AdamW + cosine + warmup."],
            ["weight_decay", "1e-4", "Chuẩn cho AdamW + Transformer."],
            ["focal_gamma", "1.0", "Focal weighting nhẹ trên class-balanced reweight."],
            ["cb_beta", "0.999", "Effective-number của Cui et al. (2019)."],
            ["aug.p_mix", "0.5", "Nửa batch là synthetic, nửa là PROVEDIt thật."],
            ["aug.max_synth_noc", "5", "PROVEDIt label tối đa 5; không sinh hơn."],
            ["aug.jitter_sigma", "0.12", "LN σ≈0.12 khớp variability re-injection."],
            ["aug.dropout_rate", "0.03", "Tỷ lệ drop-out simulation."],
            ["TTA samples", "20", "20 forward pass là sweet spot cho MC-Dropout."],
        ])

    # ==================== 9. CHẠY ====================
    add_heading(doc, "9. Cách chạy", level=1)
    add_code(doc,
        "# 1. Prepare dữ liệu (nếu chưa có .npy)\n"
        "python main.py prepare --data-dir \"data/...\"\n\n"
        "# 2. Chạy baseline để có mốc so sánh\n"
        "python main.py baseline\n\n"
        "# 3. Train deepNoC gốc với split mới (để so sánh fair)\n"
        "python main.py train --model simple --split grouped --epochs 200\n"
        "python main.py train --model full   --split grouped --epochs 200\n\n"
        "# 4. Train NoCFormer (mô hình đề xuất)\n"
        "python main.py train --model nocformer --split grouped --epochs 200 \\\n"
        "    --batch-size 32 --lr 3e-4 --tta-samples 20\n\n"
        "# 5. Smoke run nhanh (vài phút trên CPU)\n"
        "python main.py train --model nocformer --epochs 3 --batch-size 16 \\\n"
        "    --d-model 64 --peak-layers 1 --locus-layers 2 --tta-samples 4")

    # ==================== 10. KẾ HOẠCH VALIDATION ====================
    add_heading(doc, "10. Kế hoạch validation", level=1)
    add_bullet(doc, "Stage A — sanity. 2-epoch run trên CPU để verify shape, "
                    "gradient flow, augmenter, TTA. (Đã xong; 11 s/epoch ở "
                    "d_model=64, 37 s/epoch ở d_model=128.)")
    add_bullet(doc, "Stage B — head-to-head trên PROVEDIt. NoCFormer vs "
                    "deepNoC simple/full ở CÙNG grouped split, 200 epoch. "
                    "Báo cáo accuracy per-NoC, MAE, macro-F1, threshold-vs-coverage.")
    add_bullet(doc, "Stage C — ablation: (i) bật/tắt augmentation, "
                    "(ii) CORN vs CE, (iii) Transformer vs CNN, "
                    "(iv) grouped vs alternating split. Riêng việc đổi split "
                    "alone sẽ làm accuracy deepNoC giảm vài điểm — đó là leakage trước đây.")
    add_bullet(doc, "Stage D — calibration: reliability diagram (binned "
                    "confidence vs accuracy) cho softmax (deepNoC) vs "
                    "MC-Dropout mean (NoCFormer). Kỳ vọng ECE giảm 3-5 điểm.")
    add_bullet(doc, "Stage E — robustness: thêm 5-20% multiplicative height "
                    "noise khi inference; đo độ ổn định rank-1.")

    # ==================== 11. RỦI RO ====================
    add_heading(doc, "11. Rủi ro & câu hỏi mở", level=1)
    add_bullet(doc, "Synthetic mixture bỏ qua hiệu ứng stochastic (stutter "
                    "ratio biến thiên, drop-in, pull-up) mà GAN simulator "
                    "thực sự sẽ mô hình hoá. Mitigation: kết hợp với "
                    "jitter/dropout/drop-in augment, sau này có thể plug "
                    "GAN của Taylor & Humphries (2024) vào augment.py.")
    add_bullet(doc, "PROVEDIt chỉ label tới 5 người. NoC>5 là extrapolation. "
                    "CORN head support K lớn hơn nhưng cần dữ liệu giả lập.")
    add_bullet(doc, "Population mismatch (allele frequency Australian Caucasian "
                    "trong simulator gốc vs donor PROVEDIt là người Mỹ) — "
                    "DEFAULT_ALLELE_FREQ trong constants.py cần thay bằng bảng thật.")
    add_bullet(doc, "peak_label_probability vẫn là heuristic cho đến khi "
                    "tích hợp MHCNN của Taylor (2022).")

    # ==================== 12. KẾT LUẬN ====================
    add_heading(doc, "12. Kết luận", level=1)
    add_para(doc,
        "NoCFormer giữ phần đúng của deepNoC (input [24, 50, 89], multi-task "
        "explainability, GlobalFiler + PROVEDIt) và thay phần làm nó hạn chế "
        "(CNN trên trục peak tuỳ ý, softmax categorical cho label ordinal, "
        "split leak, phụ thuộc GAN, inference deterministic). Triển khai đã "
        "wired vào repo qua --model nocformer; smoke test end-to-end xác nhận "
        "shape + gradient + augment + TTA chạy. Bước kế tiếp là chạy 200-epoch "
        "head-to-head với deepNoC dưới grouped split.")

    out_path = os.path.join(here, "NoCFormer_report_VI.docx")
    doc.save(out_path)
    print(f"Báo cáo đã ghi: {out_path}")
    return out_path


if __name__ == "__main__":
    build()
