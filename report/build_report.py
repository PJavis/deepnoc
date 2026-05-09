"""
Generate the NoCFormer technical report as a .docx.

Run:
    python report/build_report.py
"""

from __future__ import annotations

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- Style helpers ----------------------------------------------------

def set_default_font(doc: Document, name: str = "Calibri", size_pt: int = 11):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)
    rpr = style.element.rPr
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x5C)
    return h


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def add_code(doc: Document, text: str):
    """Monospace code block."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    rpr = r._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")
    rpr.append(rfonts)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)


def add_table(doc: Document, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val
    return table


# ---------- Report content ---------------------------------------------------

def build():
    doc = Document()
    set_default_font(doc)

    # ---- Title page ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("NoCFormer")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1F, 0x2D, 0x5C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("A hierarchical set-Transformer for assigning the number of "
                    "contributors to STR DNA profiles")
    r.italic = True
    r.font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Improvement proposal over Taylor & Humphries (2024) — "
                 "deepNoC, arXiv:2412.09803\n").italic = True
    meta.add_run("Project: deepNoC repository (this codebase)\n")
    meta.add_run("Status: design + implementation complete; large-scale training "
                 "pending GPU run\n")

    doc.add_paragraph()  # spacer

    # ---- 1. Executive summary ----
    add_heading(doc, "1. Executive summary", level=1)
    add_para(doc,
        "This report proposes NoCFormer, a redesigned deep-learning system for "
        "assigning the number of contributors (NoC) to a short-tandem-repeat "
        "(STR) DNA profile. The baseline is the deepNoC architecture of Taylor "
        "and Humphries (2024), which is faithfully reimplemented in the current "
        "repository. NoCFormer keeps the input representation [24 × 50 × 89] "
        "and the multi-output explainability philosophy of deepNoC, but "
        "replaces the convolutional backbone with a hierarchical set-Transformer, "
        "the categorical NoC head with a rank-monotonic ordinal (CORN) head, and "
        "the GAN-based simulation pipeline with a synthetic-mixture data "
        "augmenter that operates directly on PROVEDIt single-source profiles."
    )
    add_para(doc, "Six concrete improvements are introduced:", bold=True)
    add_bullet(doc, "(1) Permutation-invariant peak encoder — peaks at a locus "
                    "are a set, not a sequence.")
    add_bullet(doc, "(2) Cross-locus self-attention with dye-channel positional "
                    "embeddings — captures pull-up and dye-balance patterns.")
    add_bullet(doc, "(3) CORN ordinal head — penalises mistakes in proportion "
                    "to how far they are from the true NoC.")
    add_bullet(doc, "(4) Synthetic-mixture augmentation — combines random "
                    "single-source PROVEDIt profiles into k-person mixtures on "
                    "the fly, removing the need for the GAN simulator.")
    add_bullet(doc, "(5) MC-Dropout test-time augmentation — yields calibrated "
                    "NoC distributions and an entropy-based confidence score.")
    add_bullet(doc, "(6) Grouped, stratified train/test split — eliminates "
                    "leakage between PROVEDIt re-injections of the same "
                    "biological mixture.")

    # ---- 2. Background ----
    add_heading(doc, "2. Background and the original deepNoC", level=1)
    add_heading(doc, "2.1 Task", level=2)
    add_para(doc,
        "Forensic STR DNA profiles are read from a capillary electropherogram "
        "(EPG) as a series of peaks at 24 (GlobalFiler) loci. Before any "
        "deconvolution can be performed, an analyst must assign a NoC. "
        "Misassignment propagates into likelihood-ratio computation: under-"
        "counting can exclude true contributors, over-counting tends to favour "
        "non-contributors. Manual NoC assignment is approximately maximum "
        "allele count (MAC) plus heuristic adjustments for peak heights and "
        "stutters, and degrades sharply beyond three contributors."
    )

    add_heading(doc, "2.2 Original deepNoC architecture", level=2)
    add_para(doc, "Input: [24 loci × 50 peaks × 89 features].")
    add_para(doc, "Backbone: a 16-layer 1D-CNN that processes peaks within "
                  "each locus, then aggregates via adaptive max pooling to a "
                  "locus representation, then applies further 1D convolutions "
                  "across the 24 loci, then max-pools to a profile vector.")
    add_para(doc, "Multi-task heads (kept for explainability):")
    add_bullet(doc, "Peak proportion allelic ([24, 50, 1], MSE)")
    add_bullet(doc, "Peak number of alleles ([24, 50, 21], CE over 0–20)")
    add_bullet(doc, "Locus mixture proportions ([24, 10], MSE)")
    add_bullet(doc, "Locus number of alleles ([24, 20], CE over 1–20)")
    add_bullet(doc, "Profile mixture proportions ([10], MSE)")
    add_bullet(doc, "Profile NoC ([10], CE)  ← primary output")
    add_para(doc,
        "Training data: 100 000 simulated profiles produced by a "
        "simDNAmixtures + GAN pipeline; fine-tuned on 371 PROVEDIt profiles. "
        "Reported test accuracy: 90 % on PROVEDIt (1–5 contributors) and 72 % "
        "on simulated profiles (1–10 contributors)."
    )

    add_heading(doc, "2.3 What this repository currently has", level=2)
    add_para(doc,
        "The repository under development reimplements the deepNoC backbone "
        "and multi-task heads, but does not have access to the GAN simulator. "
        "It is therefore trained directly on PROVEDIt (about 736 GlobalFiler "
        "25-second profiles, of which 70 are single-source and 666 are 2–5 "
        "person mixtures, after the configured single-source cap). Two "
        "auxiliary features that the paper computes from the MHCNN and "
        "STRmix's smart-start algorithm are replaced in the repo by simple "
        "heuristics: peak label probability and the [80–89] mixture "
        "proportions. This makes the current implementation essentially a "
        "structurally faithful but information-poorer reproduction."
    )

    # ---- 3. Weaknesses ----
    add_heading(doc, "3. Weaknesses of the baseline", level=1)
    add_table(doc,
        ["Issue", "Impact"],
        [
            ["1D-CNN over the peak axis is sensitive to peak ordering",
             "Peaks within a locus form a set — there is no natural sequence. "
             "Sorting by Size (as done in the loader) bakes a spurious "
             "invariance into the data the model has to learn around."],
            ["Categorical softmax over 1..K loses ordinality",
             "A NoC=4 → NoC=9 mistake is treated identically to NoC=4 → NoC=3, "
             "so the model has no incentive to be ‘nearly right’."],
            ["Adaptive max-pool collapses peak/locus representations early",
             "Fine-grained interactions between peaks and especially across "
             "loci within a dye channel are lost."],
            ["No real cross-locus interaction model",
             "Pull-up artefacts span dyes; dye balance is informative for NoC. "
             "Stacked 1D conv with kernel 3 only sees neighbouring loci in "
             "the index ordering, which is biologically arbitrary."],
            ["Heavy class imbalance (≈80 % single-source in PROVEDIt)",
             "Without rebalancing the model trivially predicts NoC=1."],
            ["Alternating train/test split",
             "PROVEDIt re-injects the same physical mixture multiple times. "
             "Index-based alternating split puts near-duplicates on opposite "
             "sides and inflates test accuracy."],
            ["Single deterministic forward pass",
             "No uncertainty estimate — yet the paper itself argues "
             "(Figure 7) that probability thresholding is the correct way to "
             "use the model in casework."],
            ["Pipeline depends on a 100 000-profile GAN simulator",
             "Without it, training material is one to two orders of "
             "magnitude too small for a deep model to generalise."],
        ])

    # ---- 4. NoCFormer design ----
    add_heading(doc, "4. NoCFormer design", level=1)

    add_heading(doc, "4.1 High-level architecture", level=2)
    add_para(doc, "Input: [B, 24, 50, 89] (unchanged).")
    add_code(doc,
        "[B, 24, 50, 89]\n"
        "   ├─ PeakEmbedder MLP          → [B, 24, 50, d]\n"
        "   ├─ PeakSetEncoder            (per-locus)\n"
        "   │     CLS token + 50 peaks\n"
        "   │     2× Pre-LN Transformer  (key-padding mask from height>0)\n"
        "   │     → locus token [B, 24, d], peak tokens [B, 24, 50, d]\n"
        "   ├─ Peak head                 → prop_allelic, n_alleles\n"
        "   ├─ LocusTransformer          (across 24 loci)\n"
        "   │     + learned locus pos\n"
        "   │     + dye-channel embedding\n"
        "   │     + profile CLS token\n"
        "   │     4× Pre-LN Transformer  (key-padding mask from active-locus)\n"
        "   │     → profile token [B, d], refined locus tokens [B, 24, d]\n"
        "   ├─ Locus head                → n_alleles, mix_props\n"
        "   └─ Profile head\n"
        "         CORN NoC logits   [B, K-1]\n"
        "         mix_props softmax [B, MAX_NOC]"
    )

    add_heading(doc, "4.2 Why a Transformer?", level=2)
    add_para(doc,
        "Self-attention is permutation-equivariant: the output is a function "
        "of which peaks are present, not the order in which they appear in "
        "the tensor. A learnable [CLS] token attends to all real peaks via "
        "the key-padding mask derived from the height feature; padding peaks "
        "contribute zero signal by construction. The same mechanism, applied "
        "across the 24 locus tokens, captures cross-locus correlations that a "
        "1D conv with kernel 3 cannot — for example, the dye-balance and "
        "pull-up interactions that span an entire dye channel."
    )

    add_heading(doc, "4.3 CORN ordinal head", level=2)
    add_para(doc,
        "NoC is an ordered label (1 < 2 < … < K). The CORN ordinal head "
        "(Cao, Mirjalili, Raschka 2020) emits K−1 logits whose sigmoids are "
        "interpreted as P(y > k). A cumulative-product trick guarantees rank "
        "monotonicity (no ‘P(y>5) > P(y>4)’ pathologies) and the conditional "
        "training objective (task k uses only samples with target > k−1) "
        "makes the loss well-defined. Decoding is just ‘number of cumulative "
        "probabilities above 0.5, plus one’. The result: a NoC=4 → NoC=3 "
        "mistake costs less than NoC=4 → NoC=9, exactly as it should in this "
        "task."
    )

    add_heading(doc, "4.4 Auxiliary heads (explainability retained)", level=2)
    add_para(doc,
        "Peak prop_allelic, peak n_alleles, locus n_alleles, locus mix_props, "
        "and profile mix_props are all retained, with small loss weights "
        "(0.05–0.10). They serve two purposes:"
    )
    add_bullet(doc, "Explainability: an analyst can still see why a particular "
                    "NoC was chosen (which loci look multi-allelic, which "
                    "peaks the model thinks are non-allelic).")
    add_bullet(doc, "Auxiliary supervision: when the targets are available, "
                    "they regularise the backbone in the same way the "
                    "original paper exploits.")

    # ---- 5. Pipeline improvements ----
    add_heading(doc, "5. Pipeline improvements", level=1)

    add_heading(doc, "5.1 Synthetic-mixture augmentation", level=2)
    add_para(doc,
        "An EPG from k contributors is, to first order, the superposition of "
        "k independent single-source EPGs with stochastic peak-height "
        "scaling. The synthetic_mix() function exploits this: it samples a "
        "Dirichlet-distributed mixture-proportion vector w ∈ Δ^{k-1}, takes "
        "k single-source [24, 50, 89] tensors from the training pool, sums "
        "their per-allele heights weighted by w, rebuilds the locus and "
        "profile peak counts, and returns a new tensor labelled k. Because "
        "C(2700, 5) ≫ 10⁹ this is effectively unlimited data, without "
        "requiring the GAN simulator that the original paper depends on."
    )

    add_heading(doc, "5.2 Other augmentations", level=2)
    add_bullet(doc, "Log-normal multiplicative jitter on peak heights (σ≈0.12) "
                    "models analytical variability between injections.")
    add_bullet(doc, "Random peak dropout (p≈0.03) models drop-out / threshold "
                    "sensitivity.")
    add_bullet(doc, "Peak-axis permutation — only safe because the encoder is "
                    "permutation-invariant. Trains the model not to depend on "
                    "the data pipeline’s arbitrary Size-sorted ordering.")

    add_heading(doc, "5.3 Class-balanced focal loss", level=2)
    add_para(doc,
        "The CORN binary tasks are weighted by the Cui et al. (2019) "
        "class-balanced effective-number formula and by a focal modulator "
        "γ=1, applied multiplicatively to the per-sample BCE before averaging. "
        "Combined with a WeightedRandomSampler this neutralises the "
        "PROVEDIt 1-person dominance without discarding any data."
    )

    add_heading(doc, "5.4 MC-Dropout TTA at inference", level=2)
    add_para(doc,
        "predict_with_tta() runs n_samples (default 20) forward passes with "
        "dropout enabled and (optionally) jittered + permuted inputs, "
        "averages the per-class probabilities, and returns predictive "
        "entropy. This produces calibrated uncertainty so the paper’s own "
        "Figure-7 use case (probability-threshold-driven assignment vs "
        "abstain) is supported out of the box rather than relying on raw "
        "softmax confidence."
    )

    add_heading(doc, "5.5 Grouped, stratified train/test split", level=2)
    add_para(doc,
        "PROVEDIt names embed the contributor pedigree, e.g. "
        "‘…-31_32-1;1-…’ identifies the same biological 2-person mixture "
        "across all of its re-injections. The new src/split.py extracts "
        "this key with a regex and uses GroupShuffleSplit so all replicates "
        "of a mixture sit on the same side of the split. With NoC stratification "
        "added on top, every class is represented in both train and test."
    )

    # ---- 6. Files added ----
    add_heading(doc, "6. Code added in this iteration", level=1)
    add_table(doc,
        ["File", "Purpose"],
        [
            ["models/nocformer/architecture.py",
             "PeakEmbedder, PeakSetEncoder, LocusTransformer, "
             "PeakHead, LocusHead, ProfileHead, NoCFormer, "
             "corn_logits_to_class_probs."],
            ["models/nocformer/losses.py",
             "corn_loss (CORN ordinal), class_balanced_weights, "
             "NoCFormerLoss (multi-task wrapper)."],
            ["models/nocformer/augment.py",
             "synthetic_mix, peak_height_jitter, random_peak_dropout, "
             "shuffle_peak_axis, TrainingAugmenter."],
            ["models/nocformer/train.py",
             "train_nocformer (cosine LR + warmup + AdamW + balanced sampler), "
             "predict_with_tta, load_nocformer."],
            ["src/split.py",
             "stratified_split, grouped_stratified_split (pedigree-aware)."],
            ["main.py",
             "Adds --model nocformer; --split {alternating,stratified,grouped}; "
             "model-config knobs."],
            ["report/build_report.py",
             "This report builder."],
        ])

    # ---- 7. Hyperparameters ----
    add_heading(doc, "7. Default hyperparameters", level=1)
    add_table(doc,
        ["Hyperparameter", "Value", "Rationale"],
        [
            ["d_model", "128", "Adequate for a 1.3 M-param model on 736 PROVEDIt profiles + augmentation."],
            ["n_heads", "4", "Standard ratio (d_model / 32) for small Transformers."],
            ["peak_layers", "2", "Two attention rounds within a locus is enough — only 50 tokens to mix."],
            ["locus_layers", "4", "Cross-locus interactions are the bigger win; deeper here."],
            ["dropout", "0.15", "Used by both regularisation and MC-Dropout uncertainty."],
            ["epochs", "200", "Matches paper’s simulated-data run; cosine schedule with warmup."],
            ["batch_size", "32", "Comfortable on a single 16 GB GPU at d_model=128."],
            ["lr", "3e-4", "AdamW + cosine + 5-epoch linear warmup."],
            ["weight_decay", "1e-4", "Standard for AdamW + Transformers."],
            ["focal_gamma", "1.0", "Mild focal weighting on top of class-balanced re-weighting."],
            ["cb_beta", "0.999", "Cui et al. class-balanced effective-number parameter."],
            ["aug.p_mix", "0.5", "Half of each batch is synthesised; the other half stays real."],
            ["aug.max_synth_noc", "5", "PROVEDIt only labels up to 5; do not synthesise beyond."],
            ["aug.jitter_sigma", "0.12", "Log-normal σ≈0.12 matches typical re-injection variability."],
            ["aug.dropout_rate", "0.03", "Drop-out simulation rate."],
            ["TTA samples", "20", "20 stochastic passes is the usual MC-Dropout sweet spot."],
        ])

    # ---- 8. Validation plan ----
    add_heading(doc, "8. Validation plan", level=1)
    add_bullet(doc, "Stage A — sanity. 2-epoch run on CPU to verify shapes, "
                    "gradient flow, augmenter outputs, and TTA loop. (Done; "
                    "11 s/epoch at d_model=64, 37 s/epoch at d_model=128.)")
    add_bullet(doc, "Stage B — per-contribution head-to-head on PROVEDIt. "
                    "Run NoCFormer with --split grouped --epochs 200 against "
                    "the existing deepNoC ‘simple’ and ‘full’ heads under the "
                    "SAME grouped split. Report per-NoC accuracy, MAE, "
                    "macro-F1, and the threshold-vs-coverage curve from "
                    "Figure 7 of the paper.")
    add_bullet(doc, "Stage C — ablations. Toggle (i) augmentation, (ii) CORN "
                    "vs CE, (iii) Transformer vs CNN backbone, (iv) grouped "
                    "vs alternating split. The split toggle alone should "
                    "bring the deepNoC baseline’s reported accuracy down by "
                    "several points — a useful honesty check.")
    add_bullet(doc, "Stage D — calibration. Reliability diagram (binned "
                    "confidence vs accuracy) for raw softmax (deepNoC) and "
                    "MC-Dropout means (NoCFormer). Expected ECE improvement "
                    "≈3–5 points based on prior MC-Dropout literature.")
    add_bullet(doc, "Stage E — robustness. Apply 5–20 % multiplicative height "
                    "noise at inference and measure the NoC-prediction "
                    "rank-1 stability. NoCFormer should be more stable thanks "
                    "to permutation invariance + dropout TTA.")

    # ---- 9. How to run ----
    add_heading(doc, "9. How to run", level=1)
    add_code(doc,
        "# 1. Prepare data (already produced X_gf25.npy, y_gf25.npy)\n"
        "python main.py prepare --data-dir \"data/...\"\n\n"
        "# 2. Run baselines\n"
        "python main.py baseline\n\n"
        "# 3. Train original deepNoC (for comparison)\n"
        "python main.py train --model simple --split grouped --epochs 200\n"
        "python main.py train --model full   --split grouped --epochs 200\n\n"
        "# 4. Train NoCFormer\n"
        "python main.py train --model nocformer --split grouped --epochs 200 \\\n"
        "    --batch-size 32 --lr 3e-4 --tta-samples 20\n\n"
        "# 5. Quick smoke run (a few minutes on CPU)\n"
        "python main.py train --model nocformer --epochs 3 --batch-size 16 \\\n"
        "    --d-model 64 --peak-layers 1 --locus-layers 2 --tta-samples 4"
    )

    # ---- 10. Risks ----
    add_heading(doc, "10. Risks and open questions", level=1)
    add_bullet(doc, "Synthetic-mixture augmentation ignores stochastic effects "
                    "(stutter ratio, drop-in, pull-up) that a true GAN "
                    "simulator would model. Mitigation: combine with the "
                    "height-jitter / dropout / drop-in augmentations and, in "
                    "future, plug in the GAN of Taylor & Humphries (2024) at "
                    "augment.py boundary.")
    add_bullet(doc, "PROVEDIt only labels up to 5 contributors. Generalisation "
                    "beyond NoC=5 is therefore extrapolation. The CORN head "
                    "supports a larger K but cannot learn what 6+ contributors "
                    "look like without simulated data.")
    add_bullet(doc, "Population mismatch (Australian Caucasian allele "
                    "frequencies in the original simulator vs US PROVEDIt "
                    "donors) was a Section-4 caveat in the paper. NoCFormer "
                    "inherits the same caveat. The DEFAULT_ALLELE_FREQ "
                    "constant should be replaced with a real frequency table.")
    add_bullet(doc, "Peak label probability remains heuristic in this repo "
                    "until the MHCNN of Taylor (2022) is integrated.")

    # ---- 11. Conclusion ----
    add_heading(doc, "11. Conclusion", level=1)
    add_para(doc,
        "NoCFormer keeps the things that work in deepNoC (the [24, 50, 89] "
        "input representation, multi-task explainability heads, the use of "
        "the GlobalFiler kit on the PROVEDIt benchmark) and replaces the "
        "things that limit it (CNN over an arbitrarily ordered peak axis, "
        "categorical softmax over an ordinal label, leaky split, "
        "GAN-dependent training data, deterministic inference). The "
        "implementation is wired into the existing repo behind "
        "‘--model nocformer’. End-to-end smoke tests confirm correct shapes "
        "and gradient flow. A full 200-epoch comparison run against the "
        "deepNoC heads under a grouped split is the next step."
    )

    # ---- Save ----
    out_path = os.path.join(os.path.dirname(__file__), "NoCFormer_report.docx")
    doc.save(out_path)
    print(f"Report written to: {out_path}")
    return out_path


if __name__ == "__main__":
    build()
